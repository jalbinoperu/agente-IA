
import streamlit as st
from docx import Document
from io import BytesIO

st.title("📘 Asistente Pedagógico IA")
st.subheader("Generador de Sesiones de Aprendizaje – Ciencia y Tecnología (Secundaria)")

# Entradas del docente
docente = st.text_input("👨‍🏫 Nombre del docente", "Jimmy Albino Meneses")
colegio = st.text_input("🏫 Nombre del colegio", "I.E. Dora Mayer")
grado = st.selectbox("📚 Grado", ["1°", "2°", "3°", "4°", "5°"])
competencia = st.selectbox("🧠 Selecciona la competencia a trabajar", [
    "Indaga mediante métodos científicos para construir conocimientos.",
    "Explica el mundo físico basándose en conocimientos sobre los seres vivos, materia y energía, biodiversidad, Tierra y universo.",
    "Diseña y construye soluciones tecnológicas para resolver problemas de su entorno."
])
titulo_sesion = st.text_input("📝 Título de la sesión", "Exploramos el método científico en nuestro entorno")

# Generación del documento
if st.button("📄 Generar sesión robusta"):
    doc = Document()
    doc.add_heading('SESIÓN DE APRENDIZAJE', 0)

    # Datos generales
    doc.add_heading('1. DATOS GENERALES', level=1)
    doc.add_paragraph(f"Docente: {docente}")
    doc.add_paragraph(f"I.E.: {colegio}")
    doc.add_paragraph(f"Área: Ciencia y Tecnología")
    doc.add_paragraph(f"Grado: {grado}")
    doc.add_paragraph(f"Duración: 90 minutos")
    doc.add_paragraph(f"Fecha: ...................")

    # Título y propósito
    doc.add_heading('2. TÍTULO DE LA SESIÓN', level=1)
    doc.add_paragraph(titulo_sesion)

    doc.add_heading('3. PROPÓSITO', level=1)
    doc.add_paragraph("Que los estudiantes comprendan, a través del trabajo colaborativo e indagación, cómo aplicar la competencia seleccionada para resolver una situación problemática cercana a su contexto.")

    # Competencia
    doc.add_heading('4. COMPETENCIA, CAPACIDAD Y DESEMPEÑO', level=1)
    doc.add_paragraph(f"Competencia: {competencia}")
    doc.add_paragraph("Capacidad: Problematiza situaciones / Diseña estrategias / Analiza datos.")
    doc.add_paragraph("Desempeño: Formula preguntas investigables, propone hipótesis, y evalúa sus resultados con base científica.")

    # Actividades de aprendizaje
    doc.add_heading('5. ACTIVIDADES DE APRENDIZAJE', level=1)
    doc.add_heading('Inicio (15 min)', level=2)
    doc.add_paragraph("- Se presenta una situación problemática real del entorno.")
- Pregunta detonante: ¿Cómo podemos demostrar que el aire ocupa espacio?
- Activación de saberes previos mediante lluvia de ideas.")

    doc.add_heading('Desarrollo (50 min)', level=2)
    doc.add_paragraph("- Los estudiantes realizan un experimento con materiales simples.
- Formulan hipótesis y registran observaciones.
- Discuten en grupos sus resultados y reflexionan sobre la validez del método.")

    doc.add_heading('Cierre (25 min)', level=2)
    doc.add_paragraph("- Socialización de hallazgos.
- Metacognición guiada con preguntas como: ¿Qué aprendiste hoy? ¿Cómo te sentiste trabajando en grupo?
- Registro en portafolio del proceso seguido.")

    # Evaluación
    doc.add_heading('6. EVALUACIÓN', level=1)
    doc.add_paragraph("Instrumento: Lista de cotejo")
    doc.add_paragraph("Criterios: Participa activamente en el experimento, formula hipótesis válidas, argumenta con base científica.")
    doc.add_paragraph("Evidencia: Registro del experimento, participación oral, reflexión escrita.")

    # Recursos
    doc.add_heading('7. RECURSOS Y MATERIALES', level=1)
    doc.add_paragraph("Materiales reciclables, agua, vasos, globos, papel, marcadores, papelotes.")

    # Observaciones
    doc.add_heading('8. OBSERVACIONES Y ADECUACIONES', level=1)
    doc.add_paragraph("Considerar apoyos visuales y actividades adaptadas para estudiantes con necesidades educativas especiales.")

    # Guardar en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("✅ Sesión pedagógica generada con éxito.")
    st.download_button("📥 Descargar sesión robusta (Word)", data=buffer, file_name="sesion_robusta.docx")
